import sys
import os
import pandas as pd
import qrcode
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QFileDialog,
    QLineEdit, QComboBox, QMessageBox, QCheckBox, QHBoxLayout
)
from PyQt6.QtGui import QFont, QIcon
from PyQt6.QtCore import Qt
import requests
from datetime import datetime


def check_for_update():
    url = "https://api.github.com/repos/ssskrbna/ARHerb/releases/latest"
    response = requests.get(url).json()
    latest_version = response.get("tag_name", "")
    current_version = "1.0"

    if latest_version > current_version:
        print("\u0414\u043e\u0441\u0442\u0443\u043f\u043d\u043e \u043e\u0431\u043d\u043e\u0432\u043b\u0435\u043d\u0438\u0435!")


def filter_nan(value):
    return value if pd.notna(value) else ''


def generate_qr_code(latitude, longitude, num):
    if latitude and longitude:
        google_maps_url = f"https://www.google.com/maps?q={latitude},{longitude}"
        qr = qrcode.make(google_maps_url)
        qr_path = f"qr_codes/qr_{num}.png"
        os.makedirs("qr_codes", exist_ok=True)
        qr.save(qr_path)
        return qr_path
    return ""


def remove_time_from_date(date_value):
    if isinstance(date_value, datetime):
        return date_value.date().strftime("%Y-%m-%d")
    elif isinstance(date_value, str):
        try:
            return datetime.strptime(date_value, "%Y-%m-%d %H:%M:%S").strftime("%Y-%m-%d")
        except ValueError:
            return date_value
    return date_value


def add_rows_to_table(dataframe, column, start_row, table, AR, herbarium_type, include_qr, column_map):
    current_row = start_row

    for _, row in dataframe.iterrows():
        get = lambda k: filter_nan(row.get(column_map.get(k, k), ''))

        family = get('family')
        species = get('species')
        familyrus = get('familyrus')
        speciesrus = get('speciesrus')
        region = get('region')
        date = remove_time_from_date(get('date'))
        point = get('point')
        habitats = get('habitats')
        leg = get('leg.')
        det = get('det.')
        num = get('num')
        latitude = get('N')
        longitude = get('E')

        qr_code_path = generate_qr_code(latitude, longitude, num) if include_qr else ""

        while len(table.rows) <= current_row:
            table.add_row()

        row_cells = table.rows[current_row].cells
        paragraph = row_cells[column].add_paragraph()
        title_run = paragraph.add_run(f"{AR}\n")
        title_run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph = row_cells[column].add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        paragraph.add_run(f"{family}\n").italic = True
        if herbarium_type == "Vascular plant" and familyrus and speciesrus:
            paragraph.add_run(f"{familyrus}\n").italic = True

        paragraph.add_run(f"{species}\n").italic = True
        if herbarium_type == "Vascular plant" and familyrus and speciesrus:
            paragraph.add_run(f"{speciesrus}\n").italic = True

        text_run = paragraph.add_run(f"""
{region}, {point}, {habitats}

{date}              leg.: {leg}
№ {num}                            det.: {det}
""")
        text_run.font.size = Pt(10)

        if include_qr and qr_code_path:
            run = paragraph.add_run()
            run.add_picture(qr_code_path, width=Inches(0.5))

        current_row += 1


class HerbariumApp(QWidget):
    def __init__(self):
        super().__init__()
        self.column_fields = {}
        self.initUI()
        self.setWindowIcon(QIcon("logo.ico"))

    def initUI(self):
        layout = QVBoxLayout()

        self.file_label = QLabel("Select Excel file:")
        layout.addWidget(self.file_label)

        self.file_button = QPushButton("Browse...")
        self.file_button.clicked.connect(self.select_file)
        layout.addWidget(self.file_button)

        self.herbarium_label = QLabel("Enter herbarium name:")
        self.herbarium_label.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        self.herbarium_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.herbarium_label)

        self.herbarium_input = QLineEdit()
        layout.addWidget(self.herbarium_input)

        self.type_label = QLabel("Specimen type (Fungi/Vascular plant/Bryophyta):")
        layout.addWidget(self.type_label)
        self.type_combo = QComboBox()
        self.type_combo.addItems(["Fungi", "Vascular plant", "Bryophyta"])
        layout.addWidget(self.type_combo)

        self.qr_checkbox = QCheckBox("Include QR Code")
        layout.addWidget(self.qr_checkbox)

        self.custom_fields_checkbox = QCheckBox("🛠 Manual column mapping")
        self.custom_fields_checkbox.stateChanged.connect(self.toggle_column_fields)
        layout.addWidget(self.custom_fields_checkbox)

        fields = ["family", "familyrus", "species", "speciesrus", "region", "point", "habitats",
                  "date", "leg.", "det.", "num", "N", "E"]
        for field in fields:
            row = QHBoxLayout()
            label = QLabel(f"{field} column:")
            line = QLineEdit()
            line.setPlaceholderText(f"Default: {field}")
            row.addWidget(label)
            row.addWidget(line)
            layout.addLayout(row)
            self.column_fields[field] = line
            line.setVisible(False)
            label.setVisible(False)
            self.column_fields[field + '_label'] = label

        self.help_button = QPushButton("Help")
        self.help_button.clicked.connect(self.show_help)
        layout.addWidget(self.help_button)

        self.submit_button = QPushButton("Generate document")
        self.submit_button.clicked.connect(self.process_file)
        layout.addWidget(self.submit_button)

        self.setLayout(layout)
        self.setWindowTitle("ARHerb")

    def toggle_column_fields(self):
        visible = self.custom_fields_checkbox.isChecked()
        for key, field in self.column_fields.items():
            if isinstance(field, QLineEdit):
                field.setVisible(visible)
            else:
                field.setVisible(visible)

    def show_help(self):
        help_text = (
            "1. Select an Excel file containing herbarium data.\n"
            "2. Enter the herbarium name.\n"
            "3. Choose the type of herbarium specimens.\n"
            "4. Optionally define custom column names.\n"
            "5. Click 'Generate' to create the labels.\n"
            "6. The output file will be saved as a Word document."
        )
        QMessageBox.information(self, "Help", help_text)

    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select file", "", "Excel Files (*.xlsx)")
        if file_path:
            self.file_label.setText(f"Selected file: {os.path.basename(file_path)}")
            self.file_path = file_path

    def get_column_mapping(self):
        return {k: self.column_fields[k].text().strip() or k for k in self.column_fields if not k.endswith('_label')}

    def process_file(self):
        if not hasattr(self, 'file_path') or not self.file_path:
            QMessageBox.warning(self, "Error", "Please select a file!")
            return

        AR = self.herbarium_input.text().strip()
        herbarium_type = self.type_combo.currentText()
        include_qr = self.qr_checkbox.isChecked()
        column_map = self.get_column_mapping() if self.custom_fields_checkbox.isChecked() else {}

        if not AR:
            QMessageBox.warning(self, "Error", "Please enter the herbarium name!")
            return

        try:
            df = pd.read_excel(self.file_path)
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            half_index = len(df) // 2
            first_half = df.iloc[:half_index]
            second_half = df.iloc[half_index:]

            add_rows_to_table(first_half, 0, 0, table, AR, herbarium_type, include_qr, column_map)
            add_rows_to_table(second_half, 1, 0, table, AR, herbarium_type, include_qr, column_map)

            doc_path = os.path.splitext(self.file_path)[0] + '_output.docx'
            doc.save(doc_path)

            QMessageBox.information(self, "Success", f"Document saved: {doc_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {e}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = HerbariumApp()
    window.show()
    sys.exit(app.exec())
